"""
Conversor de documentos (Word, Excel, TXT) para PDF
"""

import logging
from pathlib import Path
from typing import List

import openpyxl
from docx import Document
from reportlab.pdfgen import canvas
from werkzeug.datastructures import FileStorage

from ..config import Config
from ..utils import (
    save_uploaded_file,
    create_text_pdf_canvas,
    handle_page_break,
    wrap_text
)

logger = logging.getLogger(__name__)


class DocumentToPDFConverter:
    """Conversores de documentos para PDF"""
    
    @staticmethod
    def excel_to_pdf(file: FileStorage, temp_dir: str) -> List[str]:
        """
        Converte planilha Excel para PDF
        
        Args:
            file: Arquivo Excel (.xlsx)
            temp_dir: Diretório temporário
            
        Returns:
            Lista com caminho do PDF gerado
        """
        xlsx_path = save_uploaded_file(file, temp_dir)
        pdf_path = str(Path(temp_dir) / "excel_to_pdf.pdf")
        
        c, width, height, y_position = create_text_pdf_canvas(pdf_path)
        max_chars = int((width - 2 * Config.TEXT_MARGIN) / 6)
        
        try:
            workbook = openpyxl.load_workbook(str(xlsx_path), data_only=True)
            c.setFont(Config.TEXT_FONT, 10)
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # Cabeçalho da planilha
                c.drawString(Config.TEXT_MARGIN, y_position, f"--- Planilha: {sheet_name} ---")
                y_position -= 20
                
                for row in sheet.iter_rows():
                    row_data = [str(cell.value) if cell.value is not None else "" for cell in row]
                    line_text = " | ".join(row_data)
                    
                    if len(line_text) > max_chars:
                        line_text = line_text[:max_chars] + "..."
                    
                    y_position = handle_page_break(c, y_position, height)
                    if y_position == height - Config.TEXT_MARGIN:
                        c.setFont(Config.TEXT_FONT, 10)
                    
                    c.drawString(Config.TEXT_MARGIN, y_position, line_text)
                    y_position -= 15
                
                y_position -= 30
                if y_position < Config.TEXT_MARGIN and sheet_name != workbook.sheetnames[-1]:
                    c.showPage()
                    y_position = height - Config.TEXT_MARGIN
                    c.setFont(Config.TEXT_FONT, 10)
            
            logger.info(f"Excel convertido para PDF: {file.filename}")
            
        except Exception as e:
            c.drawString(Config.TEXT_MARGIN, y_position - 20, f"Erro ao ler planilha: {e}")
            logger.error(f"Erro ao converter Excel: {e}")
        finally:
            c.save()
        
        return [pdf_path]
    
    @staticmethod
    def txt_to_pdf(file: FileStorage, temp_dir: str) -> List[str]:
        """
        Converte arquivo de texto para PDF
        
        Args:
            file: Arquivo de texto (.txt)
            temp_dir: Diretório temporário
            
        Returns:
            Lista com caminho do PDF gerado
        """
        txt_path = save_uploaded_file(file, temp_dir)
        pdf_path = str(Path(temp_dir) / "text_to_pdf.pdf")
        
        c, width, height, y_position = create_text_pdf_canvas(pdf_path)
        c.setFont(Config.TEXT_FONT, Config.TEXT_FONT_SIZE)
        
        max_chars = int((width - 2 * Config.TEXT_MARGIN) / 7)
        
        try:
            with open(str(txt_path), "r", encoding="utf-8") as f:
                for line in f:
                    text_line = line.rstrip()
                    lines = wrap_text(text_line, max_chars)
                    
                    for chunk in lines:
                        y_position = handle_page_break(c, y_position, height)
                        if y_position == height - Config.TEXT_MARGIN:
                            c.setFont(Config.TEXT_FONT, Config.TEXT_FONT_SIZE)
                        
                        c.drawString(Config.TEXT_MARGIN, y_position, chunk)
                        y_position -= Config.TEXT_LINE_SPACING
            
            logger.info(f"Texto convertido para PDF: {file.filename}")
            
        except Exception as e:
            c.drawString(Config.TEXT_MARGIN, y_position - 20, f"Erro ao ler arquivo: {e}")
            logger.error(f"Erro ao converter texto: {e}")
        finally:
            c.save()
        
        return [pdf_path]
    
    @staticmethod
    def word_to_pdf(files: List[FileStorage], temp_dir: str) -> List[str]:
        """
        Converte documentos Word para PDF
        
        Args:
            files: Lista de arquivos Word (.docx)
            temp_dir: Diretório temporário
            
        Returns:
            Lista com caminho do PDF gerado
        """
        if not isinstance(files, list):
            files = [files]
        
        pdf_path = str(Path(temp_dir) / "word_to_pdf.pdf")
        c, width, height, y_position = create_text_pdf_canvas(pdf_path)
        c.setFont(Config.TEXT_FONT, 11)
        
        max_chars = int((width - 2 * Config.TEXT_MARGIN) / 6)
        
        try:
            for file_idx, file in enumerate(files):
                docx_path = save_uploaded_file(file, temp_dir)
                doc = Document(str(docx_path))
                
                # Separador entre documentos
                if file_idx > 0:
                    c.showPage()
                    y_position = height - Config.TEXT_MARGIN
                    c.setFont("Helvetica-Bold", 12)
                    c.drawString(Config.TEXT_MARGIN, y_position, "=" * 60)
                    y_position -= 20
                    c.drawString(Config.TEXT_MARGIN, y_position, f"Documento: {file.filename}")
                    y_position -= 20
                    c.drawString(Config.TEXT_MARGIN, y_position, "=" * 60)
                    y_position -= 30
                    c.setFont(Config.TEXT_FONT, 11)
                
                # Processar parágrafos
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        lines = wrap_text(paragraph.text, max_chars)
                        
                        for line in lines:
                            y_position = handle_page_break(c, y_position, height)
                            c.drawString(Config.TEXT_MARGIN, y_position, line)
                            y_position -= 20
                
                # Processar tabelas
                for table in doc.tables:
                    y_position -= 10
                    y_position = handle_page_break(c, y_position - 50, height)
                    
                    c.setFont(Config.TEXT_FONT, 9)
                    for row in table.rows:
                        row_text = " | ".join([cell.text for cell in row.cells])
                        if len(row_text) > 100:
                            row_text = row_text[:97] + "..."
                        
                        y_position = handle_page_break(c, y_position, height)
                        c.drawString(Config.TEXT_MARGIN, y_position, row_text)
                        y_position -= 15
                    
                    y_position -= 10
                    c.setFont(Config.TEXT_FONT, 11)
                
                logger.info(f"Word convertido: {file.filename}")
            
        except Exception as e:
            logger.error(f"Erro ao converter Word para PDF: {e}")
            raise RuntimeError(f"Erro ao converter Word para PDF: {e}") from e
        finally:
            c.save()
        
        return [pdf_path]
