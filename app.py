import io
import os
import shutil
import tempfile
import zipfile

import fitz  # PyMuPDF
import ghostscript
import openpyxl
from flask import (
    Flask,
    jsonify,
    render_template_string,
    request,
    send_file,
    render_template,
)
from pdf2docx import Converter
from pdf2docx.converter import ConversionException
from PIL import Image
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 100 * 1024 * 1024  # 100MB max
app.config["UPLOAD_FOLDER"] = "uploads"
app.config["OUTPUT_FOLDER"] = "outputs"

# Criar diretórios se não existirem
os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)
os.makedirs(app.config["OUTPUT_FOLDER"], exist_ok=True)

ALLOWED_EXTENSIONS = {"pdf", "docx", "txt", "xlsx", "jpg", "jpeg", "png"}


def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


@app.route("/")
def index():
    return render_template("index.html")


def excel_to_pdf(file, temp_dir):
    xlsx_path = os.path.join(temp_dir, secure_filename(file.filename))
    file.save(xlsx_path)

    pdf_path = os.path.join(temp_dir, "excel_to_pdf.pdf")
    c = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter
    y_position = height - 50

    try:
        workbook = openpyxl.load_workbook(xlsx_path)
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            c.setFont("Helvetica", 10)
            c.drawString(50, y_position, f"--- Planilha: {sheet_name} ---")
            y_position -= 20

            for row_idx, row in enumerate(sheet.iter_rows()):
                row_data = [
                    str(cell.value) if cell.value is not None else "" for cell in row
                ]
                line_text = " | ".join(row_data)

                # Simples quebra de linha para caber na página
                max_line_width = int(
                    (width - 100) / 6
                )  # Estimativa de caracteres por linha
                if len(line_text) > max_line_width:
                    # Implementação mais robusta de quebra de linha seria necessária
                    line_text = line_text[:max_line_width] + "..."

                if y_position < 50:
                    c.showPage()
                    y_position = height - 50
                    c.setFont("Helvetica", 10)  # Reset font after new page

                c.drawString(50, y_position, line_text)
                y_position -= 15  # Espaçamento menor para linhas de planilha

            y_position -= 30  # Espaçamento entre planilhas
            if (
                y_position < 50 and sheet_name != workbook.sheetnames[-1]
            ):  # Only show new page if not last sheet
                c.showPage()
                y_position = height - 50

    except Exception as e:
        # Handle potential errors with Excel files
        c.drawString(50, y_position - 20, f"Erro ao ler planilha: {e}")
        print(f"Erro ao ler planilha Excel: {e}")

    c.save()
    return [pdf_path]


def txt_to_pdf(file, temp_dir):
    txt_path = os.path.join(temp_dir, secure_filename(file.filename))
    file.save(txt_path)

    pdf_path = os.path.join(temp_dir, "text_to_pdf.pdf")
    c = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter
    y_position = height - 50

    c.setFont("Helvetica", 12)

    try:
        with open(txt_path, "r", encoding="utf-8") as f:
            for line in f:
                # Simples quebra de linha para caber na página
                text_line = line.strip()
                max_width_px = width - 100  # Margens de 50px de cada lado

                # Estimar a largura do texto para quebrar linhas
                # ReportLab não tem quebra automática de texto complexa por default
                # Esta é uma estimativa MUITO simples; para algo robusto, precisaria de TextObject
                approx_char_width_px = 7  # Média para Helvetica 12
                chars_per_line = int(max_width_px / approx_char_width_px)

                if len(text_line) > chars_per_line:
                    # Quebra simples da linha
                    chunks = [
                        text_line[i : i + chars_per_line]
                        for i in range(0, len(text_line), chars_per_line)
                    ]
                else:
                    chunks = [text_line]

                for chunk in chunks:
                    if y_position < 50:  # Margem inferior
                        c.showPage()
                        y_position = height - 50
                        c.setFont("Helvetica", 12)  # Reset font after new page

                    c.drawString(50, y_position, chunk)
                    y_position -= 15  # Espaçamento entre linhas

    except Exception as e:
        c.drawString(50, y_position - 20, f"Erro ao ler arquivo de texto: {e}")
        print(f"Erro ao ler arquivo de texto: {e}")

    c.save()
    return [pdf_path]


@app.route("/convert", methods=["POST"])
def convert():
    if "files" not in request.files:
        return jsonify({"error": "Nenhum arquivo enviado"}), 400

    files = request.files.getlist("files")
    tool = request.form.get("tool")

    if not files or files[0].filename == "":
        return jsonify({"error": "Nenhum arquivo selecionado"}), 400

    # Validação de extensão dos arquivos enviados
    for f in files:
        if not allowed_file(f.filename):
            return jsonify({"error": f"Extensão não permitida: {f.filename}"}), 400

    # Criar diretório temporário
    temp_dir = tempfile.mkdtemp()
    response = None
    try:
        if tool == "pdf-to-images":
            output_files = pdf_to_images(files[0], temp_dir)
        elif tool == "images-to-pdf":
            output_files = images_to_pdf(files, temp_dir)
        elif tool == "merge-pdf":
            output_files = merge_pdfs(files, temp_dir)
        elif tool == "split-pdf":
            output_files = split_pdf(files[0], temp_dir)
        elif tool == "compress-pdf":
            output_files = compress_pdf(files[0], temp_dir)
        elif tool == "pdf-to-pdfa":
            output_files = pdf_to_pdfa(files, temp_dir)
        elif tool == "word-to-pdf":
            output_files = word_to_pdf(files, temp_dir)
        elif tool == "excel-to-pdf":
            output_files = excel_to_pdf(files[0], temp_dir)
        elif tool == "txt-to-pdf":
            output_files = txt_to_pdf(files[0], temp_dir)
        elif tool == "pdf-to-word":
            output_files = pdf_to_word(files[0], temp_dir)
        else:
            return jsonify({"error": "Ferramenta não suportada"}), 400

        response = build_response(output_files, temp_dir)
        return response
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        # Diretório temporário limpo após preparar resposta (BytesIO) evitando remoção antecipada
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)


def pdf_to_images(file, temp_dir):
    pdf_path = os.path.join(temp_dir, secure_filename(file.filename))
    file.save(pdf_path)

    doc = fitz.open(pdf_path)
    output_files = []

    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x resolution
        img_path = os.path.join(temp_dir, f"page_{page_num + 1}.png")
        pix.save(img_path)
        output_files.append(img_path)

    doc.close()
    return output_files


def images_to_pdf(files, temp_dir):
    images = []
    for file in files:
        img_path = os.path.join(temp_dir, secure_filename(file.filename))
        file.save(img_path)
        img = Image.open(img_path)
        if img.mode != "RGB":
            img = img.convert("RGB")
        images.append(img)

    pdf_path = os.path.join(temp_dir, "images_to_pdf.pdf")
    images[0].save(pdf_path, save_all=True, append_images=images[1:])

    return [pdf_path]


def merge_pdfs(files, temp_dir):
    merged_doc = fitz.open()

    for file in files:
        pdf_path = os.path.join(temp_dir, secure_filename(file.filename))
        file.save(pdf_path)
        doc = fitz.open(pdf_path)
        merged_doc.insert_pdf(doc)
        doc.close()

    output_path = os.path.join(temp_dir, "merged.pdf")
    merged_doc.save(output_path)
    merged_doc.close()

    return [output_path]


def split_pdf(file, temp_dir):
    pdf_path = os.path.join(temp_dir, secure_filename(file.filename))
    file.save(pdf_path)

    doc = fitz.open(pdf_path)
    output_files = []

    for page_num in range(len(doc)):
        new_doc = fitz.open()
        new_doc.insert_pdf(doc, from_page=page_num, to_page=page_num)
        output_path = os.path.join(temp_dir, f"page_{page_num + 1}.pdf")
        new_doc.save(output_path)
        new_doc.close()
        output_files.append(output_path)

    doc.close()
    return output_files


def compress_pdf(file, temp_dir):
    pdf_path = os.path.join(temp_dir, secure_filename(file.filename))
    file.save(pdf_path)

    doc = fitz.open(pdf_path)
    output_path = os.path.join(temp_dir, "compressed.pdf")
    doc.save(output_path, garbage=4, deflate=True, clean=True)
    doc.close()

    return [output_path]


def pdf_to_pdfa(files, temp_dir):
    """Converte um ou mais PDFs para PDF/A-1b usando Ghostscript."""
    if not isinstance(files, list):
        files = [files]

    output_files = []

    for file in files:
        input_path = os.path.join(temp_dir, secure_filename(file.filename))
        file.save(input_path)

        base_name, _ = os.path.splitext(os.path.basename(input_path))
        output_path = os.path.join(temp_dir, f"{base_name}_pdfa.pdf")

        gs_args = [
            "gs",
            "-dPDFA=1",
            "-dBATCH",
            "-dNOPAUSE",
            "-dNOOUTERSAVE",
            "-dUseCIEColor",
            "-sProcessColorModel=DeviceRGB",
            "-sDEVICE=pdfwrite",
            "-sColorConversionStrategy=UseDeviceIndependentColor",
            "-dPDFACompatibilityPolicy=1",
            f"-sOutputFile={output_path}",
            input_path,
        ]
        gs_args = [
            arg.encode("utf-8") if isinstance(arg, str) else arg for arg in gs_args
        ]

        try:
            ghostscript.Ghostscript(*gs_args)
        except Exception as e:
            raise RuntimeError(
                f"Erro ao converter {file.filename} para PDF/A: {e}"
            ) from e

        output_files.append(output_path)

    return output_files


def word_to_pdf(files, temp_dir):
    """
    Converte um ou múltiplos arquivos DOCX para PDF
    Se houver múltiplos arquivos, mescla todos em um único PDF
    """
    from docx import Document
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    # Criar PDF de saída
    pdf_path = os.path.join(temp_dir, "word_to_pdf.pdf")
    c = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter
    y_position = height - 50

    # Se for apenas um arquivo (compatibilidade)
    if not isinstance(files, list):
        files = [files]

    # Processar cada arquivo DOCX
    for file_idx, file in enumerate(files):
        docx_path = os.path.join(temp_dir, secure_filename(file.filename))
        file.save(docx_path)

        # Lê o documento Word
        doc = Document(docx_path)

        # Adicionar separador visual (exceto no primeiro documento)
        if file_idx > 0:
            # Quebra de página
            c.showPage()
            y_position = height - 50

            # Adicionar cabeçalho com nome do arquivo
            c.setFont("Helvetica-Bold", 12)
            c.drawString(50, y_position, f"{'=' * 60}")
            y_position -= 20
            c.drawString(50, y_position, f"Documento: {file.filename}")
            y_position -= 20
            c.drawString(50, y_position, f"{'=' * 60}")
            y_position -= 30
            c.setFont("Helvetica", 11)

        # Processar parágrafos
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Quebra texto longo em múltiplas linhas
                text = paragraph.text
                max_width = width - 100

                # Estimativa simples de largura de texto
                approx_char_width = 6
                chars_per_line = int(max_width / approx_char_width)

                words = text.split()
                lines = []
                current_line = []

                for word in words:
                    if len(" ".join(current_line + [word])) <= chars_per_line:
                        current_line.append(word)
                    else:
                        if current_line:
                            lines.append(" ".join(current_line))
                            current_line = [word]
                        else:
                            lines.append(word)

                if current_line:
                    lines.append(" ".join(current_line))

                for line in lines:
                    if y_position < 50:
                        c.showPage()
                        y_position = height - 50

                    c.drawString(50, y_position, line)
                    y_position -= 20

        # Processar tabelas (se houver)
        for table in doc.tables:
            # Adicionar espaçamento antes da tabela
            y_position -= 10

            if y_position < 100:
                c.showPage()
                y_position = height - 50

            # Desenhar linhas da tabela
            c.setFont("Helvetica", 9)
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])

                # Quebrar texto da linha se necessário
                if len(row_text) > 100:
                    row_text = row_text[:97] + "..."

                if y_position < 50:
                    c.showPage()
                    y_position = height - 50

                c.drawString(50, y_position, row_text)
                y_position -= 15

            # Espaçamento após tabela
            y_position -= 10
            c.setFont("Helvetica", 11)

    c.save()
    return [pdf_path]


def pdf_to_word(file, temp_dir):
    """
    Convert PDF to Word (.docx) format.
    """
    pdf_path = os.path.join(temp_dir, secure_filename(file.filename))
    file.save(pdf_path)

    docx_filename = os.path.splitext(secure_filename(file.filename))[0] + ".docx"
    docx_path = os.path.join(temp_dir, docx_filename)

    cv = None
    try:
        cv = Converter(pdf_path)
        cv.convert(docx_path)
    except ValueError as e:
        raise RuntimeError(f"Erro no arquivo PDF: {e}") from e
    except ConversionException as e:
        raise RuntimeError(f"Erro interno na conversão: {e}") from e
    except Exception as e:
        raise RuntimeError(f"Erro ao converter {file.filename} para Word: {e}") from e
    finally:
        if cv:
            cv.close()

    return [docx_path]


def build_response(output_files, temp_dir):
    """Monta resposta enviando arquivos como attachment sem risco de remoção prematura do diretório temporário."""
    if len(output_files) == 1:
        file_path = output_files[0]
        filename = os.path.basename(file_path)
        with open(file_path, "rb") as f:
            data = f.read()
        return send_file(io.BytesIO(data), as_attachment=True, download_name=filename)
    else:
        zip_path = os.path.join(temp_dir, "converted_files.zip")
        with zipfile.ZipFile(zip_path, "w") as zipf:
            for file_path in output_files:
                zipf.write(file_path, os.path.basename(file_path))
        with open(zip_path, "rb") as f:
            data = f.read()
        return send_file(
            io.BytesIO(data), as_attachment=True, download_name="converted_files.zip"
        )


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)
